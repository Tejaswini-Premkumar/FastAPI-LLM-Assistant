from docx import Document as DocxDocument 
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings 
from langchain_community.vectorstores import Chroma 
from langchain_core.documents import Document as LangchainDocument 
import os
# Defineing the path of the document
DOCUMENT_PATH = "laptop_details_summary.docx"
# Defineing the directory for the ChromaDB persistence
CHROMA_DB_DIR = "chroma_db"

def load_document(file_path: str) -> str:
    """
    Loads text content from a .docx file.
    """
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error loading document {file_path}: {e}")
        return ""

def get_text_chunks(text: str) -> list[str]:
    """
    Splits the text into smaller, manageable chunks.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = text_splitter.split_text(text)
    return chunks

def get_embeddings_model():
    """
    Initializes and returns a Hugging Face embeddings model.
    """
    # Using sentence transformer model
    model_name = "sentence-transformers/all-MiniLM-L6-v2"
    model_kwargs = {'device': 'cpu'} 
    encode_kwargs = {'normalize_embeddings': False} # Normalizing embeddings 

    embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    return embeddings

def setup_vector_store():
    """
    Loads the document, splits it into chunks, generates embeddings,
    and sets up a ChromaDB vector store for persistence.
    """
    if os.path.exists(CHROMA_DB_DIR) and os.listdir(CHROMA_DB_DIR):
        print(f"Loading existing ChromaDB from {CHROMA_DB_DIR}")
        embeddings = get_embeddings_model()
        vector_store = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embeddings)
    else:
        print(f"Creating new ChromaDB at {CHROMA_DB_DIR}")
        document_content = load_document(DOCUMENT_PATH)
        if not document_content:
            print("Document content is empty. Vector store will not be populated.")
            return None

        text_chunks = get_text_chunks(document_content)
        if not text_chunks:
            print("No text chunks generated. Vector store will not be populated.")
            return None
        documents= [LangchainDocument(chunk) for chunk in text_chunks]
        print(f"Type of documents: {type(documents)}")
        if documents:
            print(f"Type of first element in documents: {type(documents[0])}")
            if isinstance(documents[0], LangchainDocument):
                print(f"Content of first LangchainDocument: {documents[0].page_content[:100]}...") # Print first 100 chars
        else:
            print("Documents list is empty.")

        embeddings = get_embeddings_model()
        vector_store = Chroma.from_documents(
            documents, embeddings, persist_directory=CHROMA_DB_DIR
        )
        vector_store.persist()
        print("ChromaDB created and persisted.")

    return vector_store

# Global variable to hold the vector store instance
# This will be loaded once when the application starts.
VECTOR_STORE = setup_vector_store()
